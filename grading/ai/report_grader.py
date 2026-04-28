#!/usr/bin/env python3
"""
Feedback-oriented grader for project *reports* (PDF or Word only).

Unlike ai/grader.py (compliance checking against machine-parseable requirements),
this script takes free-form *guidelines* and asks the model to give substantive
feedback. Scoring is de-emphasized; the deliverable is narrative feedback the
instructor (or student) can use to improve the work.

Usage:
  python report_grader.py --submissions /path/to/parent --guidelines /path/to/guidelines.txt

  # Guidelines can be .txt, .md, or .pdf
  python report_grader.py --submissions ./subs --guidelines ./project_guidelines.md \\
      --output_prefix report_results

The submissions directory may be structured in either of these ways (both are
supported; you can even mix them):

  * **Flat (typical for course drop folders):** .pdf and/or .docx files live
    directly in ``submissions/`` — each file is treated as one submission, keyed
    by the file name.
  * **Per-student folders:** one subfolder per student; the script walks each
    folder (recursively) and concatenates all .pdf and .docx found there.

Other top-level file types and hidden files (``.*``) are ignored. Nonstandard
.docx files (e.g. alternate OOXML namespaces) still work: the script falls back
to reading word/document.xml from the package when python-docx cannot open the
file.
"""

import argparse
import json
import os
import re
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

import pdfplumber
from openai import OpenAI
from tqdm import tqdm

try:
    import docx
except ImportError as e:  # pragma: no cover
    raise SystemExit(
        "The python-docx package is required. Install project deps, e.g. "
        "pip install python-docx"
    ) from e

MODEL = "gpt-5.4"
client = OpenAI()

# Cap extracted text per student to keep prompts bounded (adjust if needed).
MAX_GUIDELINES_CHARS = 24_000
MAX_REPORT_CHARS = 100_000

REPORT_FEEDBACK_PROMPT = """You are a thoughtful course instructor (or writing coach) reviewing a project report.

The following CONTENT is the student's report text, extracted from their submitted PDF and/or Word file(s). There may be extraction gaps (headers/footers, odd formatting); infer intent generously.

The GUIDELINES are general expectations for this assignment—not a rigid point-by-point rubric. Use them to orient your feedback, not to assign numeric penalties.

Your priorities (in order):
1. Give specific, constructive feedback: what works well, what could be stronger, and how to improve.
2. Comment on structure, clarity, argument or technical exposition, and use of evidence (figures, tables, citations) when present.
3. Relate the work to the guidelines where helpful, in plain language. Do *not* produce a long checklist of pass/fail items.
4. Be fair and professional. If a guideline area is missing or very weak, say so briefly in the feedback, but keep the focus on learning and next steps.

You MUST return STRICT JSON only (no markdown fences, no commentary outside JSON) with this exact schema:
{
  "summary": "2-4 sentence overview of the report and your read on its quality",
  "strengths": ["bullet-style strings, specific to this submission"],
  "areas_for_improvement": ["ordered by importance; actionable suggestions"],
  "guidelines_alignment": "Short paragraph: how the report aligns with the provided guidelines in spirit—what is well addressed and what could better reflect those expectations. Not a compliance table.",
  "suggested_next_steps": ["1-5 concrete next steps the author could take if revising"]
}

If the report text is empty or unusably short, set strengths and areas_for_improvement to brief notes about missing content and explain in summary."""


# ----------------------------
# FILE EXTRACTION
# ----------------------------


def extract_pdf_text(pdf_path: str) -> str:
    text = ""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += (page.extract_text() or "") + "\n"
    except Exception as e:
        print(f"PDF extraction failed: {pdf_path} ({e})")
    return text


def _xml_local_name(tag: str) -> str:
    if "}" in (tag or ""):
        return tag.rsplit("}", 1)[-1]
    return tag or ""


def _ooxml_text_from_single_xml(data: bytes) -> str:
    """
    Pull paragraph text from one Word part (document, footnotes, etc.).
    Match elements by local name (p, t) so both standard OOXML
    (schemas.openxmlformats.org/.../2006/main) and strict / alternate
    wordprocessingml namespaces (e.g. purl.oclc.org/ooxml/...) are handled.
    """
    try:
        root = ET.fromstring(data)
    except ET.ParseError:
        return ""
    paras: list[str] = []
    for p in root.iter():
        if _xml_local_name(p.tag) != "p":
            continue
        line: list[str] = []
        for t in p.iter():
            if _xml_local_name(t.tag) == "t" and t.text:
                line.append(t.text)
        s = "".join(line).strip()
        if s:
            paras.append(s)
    return "\n".join(paras)


def _extract_docx_text_ooxml(docx_path: str) -> str:
    """
    Read .docx as a zip and parse word/*.xml. Used when python-docx fails
    (e.g. nonstandard relationships) or returns no body text.
    """
    out_parts: list[str] = []
    try:
        with zipfile.ZipFile(docx_path) as zf:
            names = set(zf.namelist())
            if "word/document.xml" in names:
                with zf.open("word/document.xml") as f:
                    main = _ooxml_text_from_single_xml(f.read())
                if main.strip():
                    out_parts.append(main)
            for label, part in (
                ("Footnotes", "word/footnotes.xml"),
                ("Endnotes", "word/endnotes.xml"),
            ):
                if part in names:
                    with zf.open(part) as f:
                        t = _ooxml_text_from_single_xml(f.read())
                    if t.strip():
                        out_parts.append(f"--- {label} ---\n{t}")
            for n in sorted(names):
                m = re.match(
                    r"^word/(header|footer)(\d+)\.xml$", n, re.IGNORECASE
                )
                if m:
                    with zf.open(n) as f:
                        t = _ooxml_text_from_single_xml(f.read())
                    if t.strip():
                        label = f"{m.group(1).title()} {m.group(2)}"
                        out_parts.append(f"--- {label} ---\n{t}")
    except (zipfile.BadZipFile, KeyError, OSError):
        return ""

    return "\n\n".join(out_parts).strip()


def extract_docx_text(docx_path: str) -> str:
    text = ""
    err: Exception | None = None
    try:
        document = docx.Document(docx_path)
        parts: list[str] = []
        for p in document.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)
    except Exception as e:
        err = e

    if text.strip():
        return text

    fallback = _extract_docx_text_ooxml(docx_path)
    if fallback.strip():
        if err is not None:
            print(
                f"DOCX: python-docx could not use {docx_path} ({err}); "
                "using OOXML fallback."
            )
        return fallback

    if err is not None:
        print(f"DOCX extraction failed: {docx_path} ({err})")
    return ""


def _report_file_kind(filename: str) -> str | None:
    """
    Return 'pdf' or 'docx' based on the filename extension (case-insensitive).
    A name ending in .docx.pdf is treated as PDF so exports like 'report.docx.pdf' work.
    """
    n = filename.lower()
    if n.endswith(".pdf"):
        return "pdf"
    if n.endswith(".docx"):
        return "docx"
    return None


def collect_report_text(student_path: str) -> str:
    """
    If ``student_path`` is a single .pdf or .docx file, extract that file only.
    If it is a directory, concatenate text from all .pdf and .docx under it
    (recursively).
    """
    student_path = os.path.normpath(student_path)
    if os.path.isfile(student_path):
        base = os.path.basename(student_path)
        kind = _report_file_kind(base)
        if not kind:
            return ""
        t = (
            extract_pdf_text(student_path)
            if kind == "pdf"
            else extract_docx_text(student_path)
        )
        if t.strip():
            return f"--- File: {base} ---\n{t.strip()}"
        return ""
    chunks: list[str] = []
    for root, _, files in os.walk(student_path):
        for name in sorted(files):
            full = os.path.join(root, name)
            kind = _report_file_kind(name)
            if kind == "pdf":
                t = extract_pdf_text(full)
            elif kind == "docx":
                t = extract_docx_text(full)
            else:
                continue
            if t.strip():
                rel = os.path.relpath(full, student_path)
                chunks.append(f"--- File: {rel} ---\n{t.strip()}")
    return "\n\n".join(chunks).strip()


def list_submission_units(submissions_dir: str) -> list[tuple[str, str]]:
    """
    Return (submission_id, path) for each unit to grade.

    * Top-level .pdf / .docx: one unit per file; id is the file name.
    * Top-level subdirectories: one unit per directory; id is the directory name;
      ``collect_report_text`` scans that tree for reports.
    """
    root = os.path.abspath(os.path.normpath(submissions_dir))
    if not os.path.isdir(root):
        return []
    out: list[tuple[str, str]] = []
    try:
        names = os.listdir(root)
    except OSError:
        return []
    for name in sorted(names):
        if name.startswith("."):
            continue
        full = os.path.join(root, name)
        if os.path.isfile(full) and _report_file_kind(name):
            out.append((name, full))
        elif os.path.isdir(full):
            out.append((name, full))
    return out


def load_guidelines_text(path: str) -> str:
    p = Path(path)
    if not p.is_file():
        raise FileNotFoundError(f"Guidelines file not found: {path}")
    ext = p.suffix.lower()
    if ext == ".pdf":
        t = extract_pdf_text(str(p))
    else:
        t = p.read_text(encoding="utf-8", errors="replace")
    t = t.strip()
    if not t:
        raise ValueError(f"No text could be loaded from guidelines: {path}")
    if len(t) > MAX_GUIDELINES_CHARS:
        t = t[:MAX_GUIDELINES_CHARS] + "\n\n[Guidelines truncated for length.]"
    return t


# ----------------------------
# AI
# ----------------------------


def _extract_json_from_response(text: str):
    s = (text or "").strip()
    if s.startswith("```"):
        lines = s.split("\n")
        if lines[0].strip().startswith("```"):
            lines = lines[1:]
        if lines and lines[-1].strip() == "```":
            lines = lines[:-1]
        s = "\n".join(lines)

    def _try_parse(raw):
        try:
            return json.loads(raw)
        except json.JSONDecodeError:
            pass
        cleaned = re.sub(r"[\n\r\t]+", " ", raw)
        try:
            return json.loads(cleaned)
        except json.JSONDecodeError:
            return None

    result = _try_parse(s)
    if result is not None:
        return result
    start, end = s.find("{"), s.rfind("}")
    if start != -1 and end != -1 and end > start:
        return _try_parse(s[start : end + 1])
    return None


def _chat(messages: list) -> str:
    try:
        response = client.chat.completions.create(
            model=MODEL,
            temperature=0.2,
            messages=messages,
        )
    except Exception as e:
        if "temperature" in str(e).lower() or "unsupported" in str(e).lower():
            response = client.chat.completions.create(
                model=MODEL,
                messages=messages,
            )
        else:
            raise
    return response.choices[0].message.content


def generate_report_feedback(guidelines: str, report_text: str) -> dict:
    body = report_text
    if len(body) > MAX_REPORT_CHARS:
        body = body[:MAX_REPORT_CHARS] + "\n\n[Report text truncated for length.]"
    user = f"""{REPORT_FEEDBACK_PROMPT}

GUIDELINES (assignment expectations):
\"\"\"
{guidelines}
\"\"\"

REPORT CONTENT:
\"\"\"
{body}
\"\"\"
"""
    content = _chat(
        [
            {
                "role": "system",
                "content": "You return only valid JSON objects matching the requested schema. No markdown code blocks.",
            },
            {"role": "user", "content": user},
        ]
    )
    parsed = _extract_json_from_response(content)
    if not isinstance(parsed, dict):
        return {
            "error": "parse_failed",
            "raw_model_output": (content or "")[:8000],
            "summary": "The model did not return valid JSON. See raw_model_output in results file.",
        }
    return parsed


# ----------------------------
# MAIN
# ----------------------------


def _load_existing_results(path: Path) -> list:
    if not path.exists():
        return []
    try:
        with open(path, encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except (json.JSONDecodeError, OSError):
        return []


def main():
    parser = argparse.ArgumentParser(
        description="Generate feedback for project reports (PDF/DOCX) from free-form guidelines."
    )
    parser.add_argument(
        "--submissions",
        required=True,
        help=(
            "Submissions root: top-level .pdf/.docx each count as one submission, "
            "or one subfolder per student with reports inside (scanned recursively)"
        ),
    )
    parser.add_argument(
        "--guidelines",
        required=True,
        help="Path to guidelines (free-form rubric): .txt, .md, or .pdf",
    )
    parser.add_argument(
        "--output",
        default="",
        help="Write JSON array of {student, feedback, report_chars} to this file "
        "(default: <output_prefix>_report_feedback.json)",
    )
    parser.add_argument(
        "--output_prefix",
        default="results",
        help="Prefix for default output name if --output is omitted",
    )
    parser.add_argument(
        "--regrade-all",
        action="store_true",
        help="Regenerate feedback for all students, ignoring existing output file",
    )
    args = parser.parse_args()

    guidelines = load_guidelines_text(args.guidelines)
    out_path = Path(
        args.output
        or f"{args.output_prefix.rstrip('.json')}_report_feedback.json"
    )

    units = list_submission_units(args.submissions)

    results: list = _load_existing_results(out_path) if not args.regrade_all else []
    done = {r.get("student") for r in results if isinstance(r, dict) and r.get("student")}

    to_run = [u for u in units if u[0] not in done]
    if done and not args.regrade_all:
        print(f"Skipping {len(done)} submission(s) with existing results in {out_path}.")

    print(f"Generating feedback for {len(to_run)} submission(s).")
    for student, student_path in tqdm(to_run):
        report_text = collect_report_text(student_path)
        if not report_text:
            if os.path.isfile(student_path):
                empty_summary = (
                    "No text extracted from this PDF or DOCX. Check encryption, "
                    "scan-only PDFs, or a corrupted file."
                )
            else:
                empty_summary = (
                    "No text extracted from PDF or DOCX in this folder. Add a .pdf or "
                    ".docx report or check file encryptions/scan-only PDFs."
                )
            entry = {
                "student": student,
                "report_chars": 0,
                "feedback": {
                    "summary": empty_summary,
                    "strengths": [],
                    "areas_for_improvement": [
                        "Submit a machine-readable PDF or a .docx with selectable text so feedback can be generated."
                    ],
                    "guidelines_alignment": "N/A — no report text to review.",
                    "suggested_next_steps": [],
                },
            }
        else:
            feedback = generate_report_feedback(guidelines, report_text)
            entry = {
                "student": student,
                "report_chars": len(report_text),
                "feedback": feedback,
            }
        results.append(entry)

    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(results, f, indent=2, ensure_ascii=False)
    print(f"Wrote {len(results)} record(s) to {out_path}")


if __name__ == "__main__":
    main()
